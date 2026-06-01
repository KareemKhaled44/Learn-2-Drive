from docx import Document
from docx.shared import Inches

cases = [
    (1, "User Registration", "POST http://localhost:8000/auth/register/\nBody: {username, email, password, confirm_password, role=user}", "201 Created, user account is created"),
    (2, "Academy Registration", "POST http://localhost:8000/auth/register/academy/\nBody: {username, email, password, confirm_password, academy_name, phone, description}", "201 Created, academy registered as pending approval"),
    (3, "User Login", "POST http://localhost:8000/auth/login/\nBody: {username/email, password}", "200 OK, access and refresh tokens returned"),
    (4, "Invalid Login", "POST http://localhost:8000/auth/login/\nBody: {username/email, wrong password}", "401 Unauthorized, login rejected"),
    (5, "Get Current User Info", "GET http://localhost:8000/auth/me/\nHeader: Authorization: Bearer <token>", "200 OK, current user profile returned"),
    (6, "Create Booking", "POST http://localhost:8000/bookings/create/\nBody: {course, trainer, scheduled_date, start_time, notes}\nHeader: Authorization: Bearer <token>", "201 Created, booking created"),
    (7, "Duplicate Booking (reject)", "POST http://localhost:8000/bookings/create/\nBody: same trainer/date/time as existing booking\nHeader: Authorization: Bearer <token>", "400 Bad Request, duplicate booking rejected"),
    (8, "Check Availability", "GET http://localhost:8000/bookings/availability/?trainer=3&course=3&date=2026-05-02\nHeader: Authorization: Bearer <token>", "200 OK, available time slots returned"),
    (9, "Add Review", "POST http://localhost:8000/api/reviews/\nBody: {content_type, object_id, rating, text}\nHeader: Authorization: Bearer <token>", "201 Created, review added"),
    (10, "Academy Profile Update", "PATCH http://localhost:8000/dashboard/profile/\nBody: {name, description}\nHeader: Authorization: Bearer <academy_token>", "200 OK, profile updated"),
]

doc = Document()

doc.add_heading('Chapter 6 - System Test Cases', level=1)

table = doc.add_table(rows=1, cols=6)
table.style = 'Light List Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'No.'
hdr_cells[1].text = 'Test Case'
hdr_cells[2].text = 'Input (method & URL & sample body/headers)'
hdr_cells[3].text = 'Expected Result'
hdr_cells[4].text = 'Actual Result'
hdr_cells[5].text = 'Status'

for no, name, inp, expected in cases:
    row_cells = table.add_row().cells
    row_cells[0].text = str(no)
    row_cells[1].text = name
    row_cells[2].text = inp
    row_cells[3].text = expected
    row_cells[4].text = ''
    row_cells[5].text = ''

# Save to workspace root
output_path = 'TestCases.docx'
doc.save(output_path)
print(f"Saved {output_path}")
